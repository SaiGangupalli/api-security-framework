#!/usr/bin/env python3
"""
Jira Story to Word Document Generator with AI-powered Fraud & Security Analysis
"""

import os
import json
import base64
from datetime import datetime
from dataclasses import dataclass
from typing import List, Dict, Optional, Any
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openai
from jira import JIRA


@dataclass
class JiraStory:
    """Data class for Jira story information"""
    key: str
    summary: str
    description: str
    acceptance_criteria: str
    labels: List[str]
    attachments: List[Dict[str, Any]]
    story_type: str
    priority: str
    status: str
    assignee: str
    reporter: str
    created: str
    updated: str


class JiraConnector:
    """Handle Jira API connections and data retrieval"""
    
    def __init__(self, server_url: str, username: str, api_token: str):
        self.jira = JIRA(server=server_url, basic_auth=(username, api_token))
    
    def get_story(self, story_key: str) -> JiraStory:
        """Retrieve a single Jira story"""
        issue = self.jira.issue(story_key, expand='attachment')
        
        # Extract acceptance criteria from description or custom field
        acceptance_criteria = self._extract_acceptance_criteria(issue)
        
        # Process attachments
        attachments = []
        for attachment in issue.fields.attachment:
            attachments.append({
                'filename': attachment.filename,
                'size': attachment.size,
                'created': attachment.created,
                'content_url': attachment.content,
                'mime_type': getattr(attachment, 'mimeType', 'unknown')
            })
        
        return JiraStory(
            key=issue.key,
            summary=issue.fields.summary,
            description=issue.fields.description or "",
            acceptance_criteria=acceptance_criteria,
            labels=issue.fields.labels,
            attachments=attachments,
            story_type=str(issue.fields.issuetype),
            priority=str(issue.fields.priority),
            status=str(issue.fields.status),
            assignee=str(issue.fields.assignee) if issue.fields.assignee else "Unassigned",
            reporter=str(issue.fields.reporter),
            created=issue.fields.created,
            updated=issue.fields.updated
        )
    
    def _extract_acceptance_criteria(self, issue) -> str:
        """Extract acceptance criteria from description or custom fields"""
        # Check for custom acceptance criteria field
        if hasattr(issue.fields, 'customfield_10000'):  # Common AC field ID
            if issue.fields.customfield_10000:
                return str(issue.fields.customfield_10000)
        
        # Extract from description if it contains AC patterns
        description = issue.fields.description or ""
        ac_patterns = [
            "Acceptance Criteria:",
            "AC:",
            "Given/When/Then",
            "Scenarios:"
        ]
        
        for pattern in ac_patterns:
            if pattern.lower() in description.lower():
                parts = description.split(pattern)
                if len(parts) > 1:
                    return parts[1].strip()
        
        return ""
    
    def download_attachment(self, attachment_url: str, output_path: str):
        """Download attachment from Jira"""
        response = self.jira._session.get(attachment_url)
        if response.status_code == 200:
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return True
        return False


class AIAnalyzer:
    """Handle AI-powered fraud and security analysis"""
    
    def __init__(self, api_key: str, model: str = "gpt-4"):
        openai.api_key = api_key
        self.model = model
    
    def analyze_fraud_security_cases(self, story: JiraStory, attachment_contents: List[str] = None) -> Dict[str, Any]:
        """Analyze story for potential fraud and security cases"""
        
        # Prepare content for analysis
        content = f"""
        Story Summary: {story.summary}
        Description: {story.description}
        Acceptance Criteria: {story.acceptance_criteria}
        Labels: {', '.join(story.labels)}
        Story Type: {story.story_type}
        Priority: {story.priority}
        """
        
        if attachment_contents:
            content += f"\nAttachment Contents:\n" + "\n".join(attachment_contents)
        
        prompt = f"""
        Analyze the following Jira story for potential fraud and security considerations:

        {content}

        Please provide:
        1. FRAUD CASES - Identify potential fraud scenarios that could arise from this feature/story
        2. SECURITY CASES - Identify security vulnerabilities, threats, and mitigation strategies
        3. RISK ASSESSMENT - Rate the overall risk level (Low/Medium/High) with justification
        4. RECOMMENDATIONS - Specific security controls and fraud prevention measures
        5. COMPLIANCE CONSIDERATIONS - Any regulatory or compliance requirements

        Format your response as structured JSON with the following keys:
        - fraud_cases: Array of fraud scenarios
        - security_cases: Array of security considerations  
        - risk_assessment: Object with level and justification
        - recommendations: Array of actionable recommendations
        - compliance_considerations: Array of compliance requirements
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a cybersecurity and fraud prevention expert analyzing software requirements."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            # Parse the JSON response
            analysis_text = response.choices[0].message.content
            
            # Extract JSON from the response
            start_idx = analysis_text.find('{')
            end_idx = analysis_text.rfind('}') + 1
            
            if start_idx != -1 and end_idx != -1:
                json_str = analysis_text[start_idx:end_idx]
                return json.loads(json_str)
            else:
                # Fallback: parse the text response manually
                return self._parse_text_analysis(analysis_text)
                
        except Exception as e:
            print(f"AI analysis failed: {e}")
            return self._generate_default_analysis()
    
    def _parse_text_analysis(self, text: str) -> Dict[str, Any]:
        """Parse text-based analysis when JSON parsing fails"""
        return {
            "fraud_cases": ["Manual review required - AI parsing failed"],
            "security_cases": ["Manual review required - AI parsing failed"],
            "risk_assessment": {"level": "Medium", "justification": "Requires manual analysis"},
            "recommendations": ["Conduct manual security review"],
            "compliance_considerations": ["Review applicable regulations"]
        }
    
    def _generate_default_analysis(self) -> Dict[str, Any]:
        """Generate default analysis structure"""
        return {
            "fraud_cases": [],
            "security_cases": [],
            "risk_assessment": {"level": "Unknown", "justification": "Analysis unavailable"},
            "recommendations": [],
            "compliance_considerations": []
        }


class WordDocumentGenerator:
    """Generate Word documents from Jira stories and AI analysis"""
    
    def __init__(self):
        self.doc = None
    
    def create_document(self, story: JiraStory, ai_analysis: Dict[str, Any], output_path: str):
        """Create a comprehensive Word document"""
        self.doc = Document()
        
        # Set up document styles
        self._setup_styles()
        
        # Document header
        self._add_header(story)
        
        # Story details section
        self._add_story_details(story)
        
        # Fraud and Security Analysis section
        self._add_fraud_security_analysis(ai_analysis)
        
        # Attachments section
        if story.attachments:
            self._add_attachments_section(story.attachments)
        
        # Footer
        self._add_footer()
        
        # Save document
        self.doc.save(output_path)
        print(f"Document saved: {output_path}")
    
    def _setup_styles(self):
        """Set up document styles"""
        styles = self.doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Heading styles
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Subheading style
        subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.size = Pt(12)
        subheading_style.font.bold = True
    
    def _add_header(self, story: JiraStory):
        """Add document header"""
        # Title
        title = self.doc.add_paragraph(f"Jira Story Analysis: {story.key}", style='CustomTitle')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Subtitle
        subtitle = self.doc.add_paragraph(story.summary, style='CustomHeading')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        self.doc.add_paragraph()  # Space
    
    def _add_story_details(self, story: JiraStory):
        """Add story details section"""
        self.doc.add_paragraph("STORY DETAILS", style='CustomHeading')
        
        # Create table for story details
        table = self.doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'
        
        # Populate table
        details = [
            ("Story Key", story.key),
            ("Type", story.story_type),
            ("Priority", story.priority),
            ("Status", story.status),
            ("Assignee", story.assignee),
            ("Reporter", story.reporter),
            ("Created", story.created),
            ("Updated", story.updated),
            ("Labels", ", ".join(story.labels) if story.labels else "None")
        ]
        
        for i, (key, value) in enumerate(details):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        self.doc.add_paragraph()
        
        # Description
        self.doc.add_paragraph("Description", style='CustomSubheading')
        self.doc.add_paragraph(story.description or "No description provided")
        
        self.doc.add_paragraph()
        
        # Acceptance Criteria
        self.doc.add_paragraph("Acceptance Criteria", style='CustomSubheading')
        self.doc.add_paragraph(story.acceptance_criteria or "No acceptance criteria provided")
        
        self.doc.add_paragraph()
    
    def _add_fraud_security_analysis(self, analysis: Dict[str, Any]):
        """Add fraud and security analysis section"""
        self.doc.add_paragraph("FRAUD & SECURITY ANALYSIS", style='CustomHeading')
        
        # Risk Assessment
        risk = analysis.get('risk_assessment', {})
        self.doc.add_paragraph("Risk Assessment", style='CustomSubheading')
        risk_para = self.doc.add_paragraph()
        risk_para.add_run(f"Risk Level: ").bold = True
        risk_para.add_run(risk.get('level', 'Unknown'))
        self.doc.add_paragraph(f"Justification: {risk.get('justification', 'Not provided')}")
        
        self.doc.add_paragraph()
        
        # Fraud Cases
        self.doc.add_paragraph("Potential Fraud Cases", style='CustomSubheading')
        fraud_cases = analysis.get('fraud_cases', [])
        if fraud_cases:
            for i, case in enumerate(fraud_cases, 1):
                self.doc.add_paragraph(f"{i}. {case}")
        else:
            self.doc.add_paragraph("No specific fraud cases identified")
        
        self.doc.add_paragraph()
        
        # Security Cases
        self.doc.add_paragraph("Security Considerations", style='CustomSubheading')
        security_cases = analysis.get('security_cases', [])
        if security_cases:
            for i, case in enumerate(security_cases, 1):
                self.doc.add_paragraph(f"{i}. {case}")
        else:
            self.doc.add_paragraph("No specific security considerations identified")
        
        self.doc.add_paragraph()
        
        # Recommendations
        self.doc.add_paragraph("Recommendations", style='CustomSubheading')
        recommendations = analysis.get('recommendations', [])
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                self.doc.add_paragraph(f"{i}. {rec}")
        else:
            self.doc.add_paragraph("No specific recommendations provided")
        
        self.doc.add_paragraph()
        
        # Compliance Considerations
        self.doc.add_paragraph("Compliance Considerations", style='CustomSubheading')
        compliance = analysis.get('compliance_considerations', [])
        if compliance:
            for i, comp in enumerate(compliance, 1):
                self.doc.add_paragraph(f"{i}. {comp}")
        else:
            self.doc.add_paragraph("No specific compliance requirements identified")
        
        self.doc.add_paragraph()
    
    def _add_attachments_section(self, attachments: List[Dict[str, Any]]):
        """Add attachments section"""
        self.doc.add_paragraph("ATTACHMENTS", style='CustomHeading')
        
        if attachments:
            for attachment in attachments:
                para = self.doc.add_paragraph()
                para.add_run(f"• {attachment['filename']}").bold = True
                para.add_run(f" ({attachment['size']} bytes)")
                self.doc.add_paragraph(f"  Type: {attachment.get('mime_type', 'Unknown')}")
                self.doc.add_paragraph(f"  Created: {attachment['created']}")
        else:
            self.doc.add_paragraph("No attachments found")
        
        self.doc.add_paragraph()
    
    def _add_footer(self):
        """Add document footer"""
        self.doc.add_paragraph()
        footer = self.doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer.runs[0].italic = True


class JiraWordGenerator:
    """Main orchestrator class"""
    
    def __init__(self, config: Dict[str, str]):
        self.jira_connector = JiraConnector(
            config['jira_server'],
            config['jira_username'],
            config['jira_api_token']
        )
        self.ai_analyzer = AIAnalyzer(config['openai_api_key'])
        self.word_generator = WordDocumentGenerator()
        self.output_dir = config.get('output_dir', 'output')
        self.attachments_dir = os.path.join(self.output_dir, 'attachments')
        
        # Create directories
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.attachments_dir, exist_ok=True)
    
    def process_story(self, story_key: str) -> str:
        """Process a single Jira story and generate Word document"""
        print(f"Processing story: {story_key}")
        
        # Get story from Jira
        story = self.jira_connector.get_story(story_key)
        
        # Download and process attachments
        attachment_contents = []
        for attachment in story.attachments:
            attachment_path = os.path.join(self.attachments_dir, f"{story_key}_{attachment['filename']}")
            if self.jira_connector.download_attachment(attachment['content_url'], attachment_path):
                print(f"Downloaded attachment: {attachment['filename']}")
                
                # Try to read text content for AI analysis
                if attachment.get('mime_type', '').startswith('text/'):
                    try:
                        with open(attachment_path, 'r', encoding='utf-8') as f:
                            attachment_contents.append(f.read())
                    except Exception as e:
                        print(f"Could not read attachment {attachment['filename']}: {e}")
        
        # Perform AI analysis
        print("Performing fraud and security analysis...")
        ai_analysis = self.ai_analyzer.analyze_fraud_security_cases(story, attachment_contents)
        
        # Generate Word document
        output_path = os.path.join(self.output_dir, f"{story_key}_analysis.docx")
        self.word_generator.create_document(story, ai_analysis, output_path)
        
        return output_path
    
    def process_multiple_stories(self, story_keys: List[str]) -> List[str]:
        """Process multiple Jira stories"""
        results = []
        for story_key in story_keys:
            try:
                result = self.process_story(story_key)
                results.append(result)
            except Exception as e:
                print(f"Error processing {story_key}: {e}")
                results.append(None)
        return results


def main():
    """Main function to demonstrate usage"""
    # Configuration
    config = {
        'jira_server': 'https://your-company.atlassian.net',
        'jira_username': 'your-email@company.com',
        'jira_api_token': 'your-jira-api-token',
        'openai_api_key': 'your-openai-api-key',
        'output_dir': 'output'
    }
    
    # Initialize generator
    generator = JiraWordGenerator(config)
    
    # Process single story
    story_key = "PROJ-123"
    output_path = generator.process_story(story_key)
    print(f"Generated document: {output_path}")
    
    # Process multiple stories
    story_keys = ["PROJ-123", "PROJ-124", "PROJ-125"]
    results = generator.process_multiple_stories(story_keys)
    print(f"Processed {len([r for r in results if r])} stories successfully")


if __name__ == "__main__":
    main()








# Jira Story to Word Document Generator - Project Structure

## Folder Structure
```
jira-word-generator/
├── README.md
├── requirements.txt
├── setup.py
├── config/
│   ├── __init__.py
│   ├── settings.py
│   └── config.yaml.example
├── src/
│   ├── __init__.py
│   ├── main.py
│   ├── jira_connector.py
│   ├── ai_analyzer.py
│   ├── word_generator.py
│   └── utils/
│       ├── __init__.py
│       ├── file_handler.py
│       └── validators.py
├── templates/
│   ├── word_template.docx
│   └── analysis_template.json
├── output/
│   ├── documents/
│   └── attachments/
├── tests/
│   ├── __init__.py
│   ├── test_jira_connector.py
│   ├── test_ai_analyzer.py
│   ├── test_word_generator.py
│   └── fixtures/
│       ├── sample_story.json
│       └── sample_analysis.json
├── scripts/
│   ├── batch_process.py
│   ├── setup_env.py
│   └── validate_config.py
├── docs/
│   ├── installation.md
│   ├── configuration.md
│   ├── usage.md
│   └── api_reference.md
└── logs/
    └── .gitkeep
```

## Requirements File (requirements.txt)
```
# Core dependencies
jira==3.5.0
python-docx==0.8.11
openai==1.3.0
requests==2.31.0
pydantic==2.5.0
python-dotenv==1.0.0
pyyaml==6.0.1

# Optional dependencies for enhanced functionality
Pillow==10.1.0  # For image processing
pandas==2.1.4   # For data analysis
openpyxl==3.1.2  # For Excel file processing
python-magic==0.4.27  # For file type detection

# Development dependencies
pytest==7.4.3
pytest-cov==4.1.0
black==23.11.0
flake8==6.1.0
mypy==1.7.1

# Logging and monitoring
structlog==23.2.0
sentry-sdk==1.38.0
```

## Setup Configuration (setup.py)
```python
from setuptools import setup, find_packages

with open("README.md", "r", encoding="utf-8") as fh:
    long_description = fh.read()

with open("requirements.txt", "r", encoding="utf-8") as fh:
    requirements = [line.strip() for line in fh if line.strip() and not line.startswith("#")]

setup(
    name="jira-word-generator",
    version="1.0.0",
    author="Your Name",
    author_email="your.email@company.com",
    description="Generate Word documents from Jira stories with AI-powered fraud and security analysis",
    long_description=long_description,
    long_description_content_type="text/markdown",
    url="https://github.com/yourcompany/jira-word-generator",
    packages=find_packages(where="src"),
    package_dir={"": "src"},
    classifiers=[
        "Development Status :: 4 - Beta",
        "Intended Audience :: Developers",
        "License :: OSI Approved :: MIT License",
        "Operating System :: OS Independent",
        "Programming Language :: Python :: 3",
        "Programming Language :: Python :: 3.8",
        "Programming Language :: Python :: 3.9",
        "Programming Language :: Python :: 3.10",
        "Programming Language :: Python :: 3.11",
    ],
    python_requires=">=3.8",
    install_requires=requirements,
    entry_points={
        "console_scripts": [
            "jira-word-gen=main:main",
        ],
    },
)
```

## Configuration File (config/settings.py)
```python
import os
import yaml
from typing import Dict, Optional
from pydantic import BaseSettings, validator
from dotenv import load_dotenv

load_dotenv()

class Settings(BaseSettings):
    # Jira Configuration
    jira_server: str
    jira_username: str
    jira_api_token: str
    
    # OpenAI Configuration
    openai_api_key: str
    openai_model: str = "gpt-4"
    
    # Output Configuration
    output_dir: str = "output"
    attachments_dir: str = "output/attachments"
    
    # Document Configuration
    include_attachments: bool = True
    max_attachment_size: int = 10 * 1024 * 1024  # 10MB
    supported_attachment_types: list = [
        "text/plain", "text/csv", "application/json",
        "application/pdf", "image/png", "image/jpeg"
    ]
    
    # AI Analysis Configuration
    enable_ai_analysis: bool = True
    ai_temperature: float = 0.3
    max_tokens: int = 2000
    
    # Logging Configuration
    log_level: str = "INFO"
    log_file: str = "logs/app.log"
    
    @validator('jira_server')
    def validate_jira_server(cls, v):
        if not v.startswith(('http://', 'https://')):
            raise ValueError('Jira server must start with http:// or https://')
        return v
    
    @validator('output_dir')
    def validate_output_dir(cls, v):
        os.makedirs(v, exist_ok=True)
        return v
    
    class Config:
        env_file = ".env"
        case_sensitive = False

def load_config(config_path: Optional[str] = None) -> Settings:
    """Load configuration from file or environment variables"""
    if config_path and os.path.exists(config_path):
        with open(config_path, 'r') as f:
            config_data = yaml.safe_load(f)
        return Settings(**config_data)
    return Settings()
```

## Environment Variables Template (.env.example)
```env
# Jira Configuration
JIRA_SERVER=https://your-company.atlassian.net
JIRA_USERNAME=your-email@company.com
JIRA_API_TOKEN=your-jira-api-token

# OpenAI Configuration
OPENAI_API_KEY=your-openai-api-key
OPENAI_MODEL=gpt-4

# Output Configuration
OUTPUT_DIR=output
ATTACHMENTS_DIR=output/attachments

# AI Analysis Configuration
ENABLE_AI_ANALYSIS=true
AI_TEMPERATURE=0.3
MAX_TOKENS=2000

# Logging Configuration
LOG_LEVEL=INFO
LOG_FILE=logs/app.log
```

## Configuration YAML Template (config/config.yaml.example)
```yaml
jira:
  server: "https://your-company.atlassian.net"
  username: "your-email@company.com"
  api_token: "your-jira-api-token"

openai:
  api_key: "your-openai-api-key"
  model: "gpt-4"
  temperature: 0.3
  max_tokens: 2000

output:
  dir: "output"
  attachments_dir: "output/attachments"
  include_attachments: true
  max_attachment_size: 10485760  # 10MB

document:
  template_path: "templates/word_template.docx"
  include_metadata: true
  include_analysis: true

security_analysis:
  enabled: true
  fraud_detection: true
  vulnerability_assessment: true
  compliance_check: true
  risk_categories:
    - "Authentication"
    - "Authorization"
    - "Data Protection"
    - "Input Validation"
    - "Session Management"
    - "Cryptography"
    - "Error Handling"
    - "Logging"

logging:
  level: "INFO"
  file: "logs/app.log"
  format: "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
```

## Batch Processing Script (scripts/batch_process.py)
```python
#!/usr/bin/env python3
"""
Batch processing script for multiple Jira stories
"""

import sys
import argparse
import csv
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent / "src"))

from main import JiraWordGenerator
from config.settings import load_config

def process_from_csv(csv_file: str, config_file: str = None):
    """Process stories from CSV file"""
    config = load_config(config_file)
    generator = JiraWordGenerator(config.dict())
    
    results = []
    with open(csv_file, 'r') as f:
        reader = csv.reader(f)
        story_keys = [row[0] for row in reader if row]
    
    print(f"Processing {len(story_keys)} stories...")
    
    for story_key in story_keys:
        try:
            output_path = generator.process_story(story_key)
            results.append({"story_key": story_key, "status": "Success", "output": output_path})
            print(f"✓ {story_key} - Generated: {output_path}")
        except Exception as e:
            results.append({"story_key": story_key, "status": "Failed", "error": str(e)})
            print(f"✗ {story_key} - Error: {e}")
    
    # Generate summary report
    success_count = len([r for r in results if r["status"] == "Success"])
    print(f"\nSummary: {success_count}/{len(story_keys)} stories processed successfully")
    
    return results

def main():
    parser = argparse.ArgumentParser(description="Batch process Jira stories")
    parser.add_argument("csv_file", help="CSV file containing story keys")
    parser.add_argument("--config", help="Configuration file path")
    
    args = parser.parse_args()
    process_from_csv(args.csv_file, args.config)

if __name__ == "__main__":
    main()
```

## Installation Script (scripts/setup_env.py)
```python
#!/usr/bin/env python3
"""
Environment setup script
"""

import os
import subprocess
import sys
from pathlib import Path

def setup_environment():
    """Set up the development environment"""
    print("Setting up Jira Word Generator environment...")
    
    # Create virtual environment
    print("Creating virtual environment...")
    subprocess.run([sys.executable, "-m", "venv", "venv"])
    
    # Determine activation script path
    if os.name == 'nt':  # Windows
        activate_script = "venv\\Scripts\\activate.bat"
        pip_path = "venv\\Scripts\\pip"
    else:  # Unix/Linux/macOS
        activate_script = "venv/bin/activate"
        pip_path = "venv/bin/pip"
    
    # Install requirements
    print("Installing requirements...")
    subprocess.run([pip_path, "install", "-r", "requirements.txt"])
    
    # Create necessary directories
    directories = [
        "output", "output/documents", "output/attachments",
        "logs", "templates", "tests/fixtures"
    ]
    
    for directory in directories:
        Path(directory).mkdir(parents=True, exist_ok=True)
        print(f"Created directory: {directory}")
    
    # Copy example configuration files
    if not os.path.exists(".env"):
        if os.path.exists(".env.example"):
            import shutil
            shutil.copy(".env.example", ".env")
            print("Created .env file from template")
    
    print("\nSetup complete!")
    print(f"To activate the virtual environment, run: {activate_script}")
    print("Don't forget to update the .env file with your credentials!")

if __name__ == "__main__":
    setup_environment()
```

## Validation Script (scripts/validate_config.py)
```python
#!/usr/bin/env python3
"""
Configuration validation script
"""

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent / "src"))

from config.settings import load_config
from jira import JIRA
import openai

def validate_jira_connection(config):
    """Validate Jira connection"""
    try:
        jira = JIRA(
            server=config.jira_server,
            basic_auth=(config.jira_username, config.jira_api_token)
        )
        # Test connection by getting server info
        server_info = jira.server_info()
        print(f"✓ Jira connection successful - Server: {server_info.get('serverTitle', 'Unknown')}")
        return True
    except Exception as e:
        print(f"✗ Jira connection failed: {e}")
        return False

def validate_openai_connection(config):
    """Validate OpenAI connection"""
    try:
        openai.api_key = config.openai_api_key
        # Test with a simple completion
        response = openai.ChatCompletion.create(
            model=config.openai_model,
            messages=[{"role": "user", "content": "Test"}],
            max_tokens=5
        )
        print("✓ OpenAI connection successful")
        return True
    except Exception as e:
        print(f"✗ OpenAI connection failed: {e}")
        return False

def validate_directories(config):
    """Validate directory structure"""
    directories = [config.output_dir, config.attachments_dir]
    all_valid = True
    
    for directory in directories:
        if Path(directory).exists():
            print(f"✓ Directory exists: {directory}")
        else:
            print(f"✗ Directory missing: {directory}")
            all_valid = False
    
    return all_valid

def main():
    try:
        config = load_config()
        print("Validating configuration...\n")
        
        # Validate connections
        jira_valid = validate_jira_connection(config)
        openai_valid = validate_openai_connection(config)
        dirs_valid = validate_directories(config)
        
        if all([jira_valid, openai_valid, dirs_valid]):
            print("\n✓ All validations passed! Configuration is ready.")
            return 0
        else:
            print("\n✗ Some validations failed. Please check your configuration.")
            return 1
            
    except Exception as e:
        print(f"Configuration validation failed: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(main())
```

## Enhanced AI Analyzer (src/ai_analyzer.py)
```python
"""
Enhanced AI Analyzer with specialized fraud and security analysis
"""

import json
import re
from typing import Dict, List, Any, Optional
import openai
from dataclasses import dataclass

@dataclass
class SecurityPattern:
    """Security pattern definitions"""
    name: str
    keywords: List[str]
    risk_level: str
    description: str

class EnhancedAIAnalyzer:
    """Enhanced AI analyzer with security pattern matching"""
    
    def __init__(self, api_key: str, model: str = "gpt-4"):
        openai.api_key = api_key
        self.model = model
        self.security_patterns = self._load_security_patterns()
        self.fraud_patterns = self._load_fraud_patterns()
    
    def _load_security_patterns(self) -> List[SecurityPattern]:
        """Load predefined security patterns"""
        return [
            SecurityPattern(
                name="Authentication Bypass",
                keywords=["login", "auth", "token", "session", "bypass", "skip"],
                risk_level="High",
                description="Potential authentication bypass vulnerabilities"
            ),
            SecurityPattern(
                name="SQL Injection",
                keywords=["database", "query", "sql", "input", "parameter"],
                risk_level="High",
                description="SQL injection attack vectors"
            ),
            SecurityPattern(
                name="XSS Vulnerability",
                keywords=["script", "html", "input", "output", "render"],
                risk_level="Medium",
                description="Cross-site scripting vulnerabilities"
            ),
            SecurityPattern(
                name="Data Exposure",
                keywords=["personal", "sensitive", "pii", "data", "information"],
                risk_level="High",
                description="Potential data exposure risks"
            ),
            SecurityPattern(
                name="Privilege Escalation",
                keywords=["admin", "privilege", "role", "permission", "access"],
                risk_level="High",
                description="Privilege escalation vulnerabilities"
            )
        ]
    
    def _load_fraud_patterns(self) -> List[SecurityPattern]:
        """Load predefined fraud patterns"""
        return [
            SecurityPattern(
                name="Payment Fraud",
                keywords=["payment", "transaction", "money", "financial", "billing"],
                risk_level="High",
                description="Payment and financial fraud risks"
            ),
            SecurityPattern(
                name="Identity Fraud",
                keywords=["identity", "profile", "account", "user", "impersonation"],
                risk_level="High",
                description="Identity theft and impersonation risks"
            ),
            SecurityPattern(
                name="Data Manipulation",
                keywords=["modify", "update", "delete", "change", "manipulate"],
                risk_level="Medium",
                description="Unauthorized data manipulation"
            ),
            SecurityPattern(
                name="Social Engineering",
                keywords=["email", "notification", "alert", "message", "communication"],
                risk_level="Medium",
                description="Social engineering attack vectors"
            )
        ]
    
    def analyze_patterns(self, content: str) -> Dict[str, List[str]]:
        """Analyze content for security and fraud patterns"""
        content_lower = content.lower()
        
        matched_security = []
        matched_fraud = []
        
        for pattern in self.security_patterns:
            if any(keyword in content_lower for keyword in pattern.keywords):
                matched_security.append(f"{pattern.name} - {pattern.description}")
        
        for pattern in self.fraud_patterns:
            if any(keyword in content_lower for keyword in pattern.keywords):
                matched_fraud.append(f"{pattern.name} - {pattern.description}")
        
        return {
            "security_patterns": matched_security,
            "fraud_patterns": matched_fraud
        }
    
    def analyze_fraud_security_cases(self, story, attachment_contents: List[str] = None) -> Dict[str, Any]:
        """Enhanced fraud and security analysis"""
        
        # Prepare content
        content = f"""
        Story Summary: {story.summary}
        Description: {story.description}
        Acceptance Criteria: {story.acceptance_criteria}
        Labels: {', '.join(story.labels)}
        Story Type: {story.story_type}
        Priority: {story.priority}
        """
        
        if attachment_contents:
            content += f"\nAttachment Contents:\n" + "\n".join(attachment_contents)
        
        # Pattern matching analysis
        pattern_analysis = self.analyze_patterns(content)
        
        # AI-powered analysis
        ai_analysis = self._perform_ai_analysis(content, pattern_analysis)
        
        # Combine results
        return self._combine_analyses(pattern_analysis, ai_analysis)
    
    def _perform_ai_analysis(self, content: str, pattern_analysis: Dict[str, List[str]]) -> Dict[str, Any]:
        """Perform AI analysis with context from pattern matching"""
        
        prompt = f"""
        As a cybersecurity and fraud prevention expert, analyze this Jira story for security and fraud risks.
        
        Story Content:
        {content}
        
        Pattern Analysis Results:
        Security Patterns Found: {pattern_analysis.get('security_patterns', [])}
        Fraud Patterns Found: {pattern_analysis.get('fraud_patterns', [])}
        
        Provide a comprehensive analysis covering:
        
        1. FRAUD SCENARIOS - Specific fraud scenarios considering the functionality
        2. SECURITY VULNERABILITIES - Technical security vulnerabilities and attack vectors
        3. RISK ASSESSMENT - Overall risk level (Critical/High/Medium/Low) with detailed justification
        4. MITIGATION STRATEGIES - Specific controls and countermeasures
        5. TESTING RECOMMENDATIONS - Security testing approaches
        6. COMPLIANCE IMPACT - Regulatory compliance considerations (GDPR, PCI-DSS, SOX, etc.)
        7. BUSINESS IMPACT - Potential business impact of identified risks
        
        Return response as valid JSON with these exact keys:
        {
            "fraud_scenarios": [],
            "security_vulnerabilities": [],
            "risk_assessment": {"level": "", "score": 0, "justification": ""},
            "mitigation_strategies": [],
            "testing_recommendations": [],
            "compliance_impact": [],
            "business_impact": ""
        }
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a senior cybersecurity analyst specializing in application security and fraud prevention. Provide detailed, actionable analysis in JSON format."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=2000
            )
            
            response_text = response.choices[0].message.content.strip()
            
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return self._generate_fallback_analysis()
                
        except Exception as e:
            print(f"AI analysis error: {e}")
            return self._generate_fallback_analysis()
    
    def _combine_analyses(self, pattern_analysis: Dict[str, List[str]], ai_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Combine pattern matching and AI analysis results"""
        
        # Calculate enhanced risk score
        risk_score = self._calculate_risk_score(pattern_analysis, ai_analysis)
        
        return {
            "fraud_scenarios": ai_analysis.get("fraud_scenarios", []),
            "security_vulnerabilities": ai_analysis.get("security_vulnerabilities", []),
            "risk_assessment": {
                "level": ai_analysis.get("risk_assessment", {}).get("level", "Medium"),
                "score": risk_score,
                "justification": ai_analysis.get("risk_assessment", {}).get("justification", ""),
                "pattern_matches": pattern_analysis
            },
            "mitigation_strategies": ai_analysis.get("mitigation_strategies", []),
            "testing_recommendations": ai_analysis.get("testing_recommendations", []),
            "compliance_impact": ai_analysis.get("compliance_impact", []),
            "business_impact": ai_analysis.get("business_impact", ""),
            "analysis_metadata": {
                "model_used": self.model,
                "pattern_matches_count": len(pattern_analysis.get("security_patterns", [])) + len(pattern_analysis.get("fraud_patterns", [])),
                "analysis_timestamp": self._get_timestamp()
            }
        }
    
    def _calculate_risk_score(self, pattern_analysis: Dict[str, List[str]], ai_analysis: Dict[str, Any]) -> int:
        """Calculate numerical risk score (0-100)"""
        base_score = 20  # Base risk score
        
        # Add points for pattern matches
        security_patterns = len(pattern_analysis.get("security_patterns", []))
        fraud_patterns = len(pattern_analysis.get("fraud_patterns", []))
        
        pattern_score = min(30, (security_patterns + fraud_patterns) * 5)
        
        # Add points based on AI risk level
        ai_risk_level = ai_analysis.get("risk_assessment", {}).get("level", "Medium")
        risk_multipliers = {"Critical": 50, "High": 40, "Medium": 25, "Low": 10}
        ai_score = risk_multipliers.get(ai_risk_level, 25)
        
        total_score = min(100, base_score + pattern_score + ai_score)
        return total_score
    
    def _generate_fallback_analysis(self) -> Dict[str, Any]:
        """Generate fallback analysis when AI fails"""
        return {
            "fraud_scenarios": ["Manual review required - AI analysis unavailable"],
            "security_vulnerabilities": ["Manual security assessment needed"],
            "risk_assessment": {"level": "Medium", "score": 50, "justification": "Default assessment - requires manual review"},
            "mitigation_strategies": ["Conduct manual security review"],
            "testing_recommendations": ["Perform standard security testing"],
            "compliance_impact": ["Review applicable compliance requirements"],
            "business_impact": "Requires manual assessment"
        }
    
    def _get_timestamp(self) -> str:
        """Get current timestamp"""
        from datetime import datetime
        return datetime.now().isoformat()
```

## Enhanced Word Generator (src/word_generator.py)
```python
"""
Enhanced Word document generator with improved formatting and analysis presentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import matplotlib.pyplot as plt
import io
import base64
from datetime import datetime
from typing import Dict, Any, List

class EnhancedWordGenerator:
    """Enhanced Word document generator with rich formatting"""
    
    def __init__(self):
        self.doc = None
        self.styles_created = False
    
    def create_document(self, story, ai_analysis: Dict[str, Any], output_path: str):
        """Create enhanced Word document"""
        self.doc = Document()
        self._setup_enhanced_styles()
        
        # Document sections
        self._add_cover_page(story)
        self._add_executive_summary(story, ai_analysis)
        self._add_story_details(story)
        self._add_risk_analysis(ai_analysis)
        self._add_detailed_findings(ai_analysis)
        self._add_recommendations(ai_analysis)
        self._add_compliance_section(ai_analysis)
        
        if story.attachments:
            self._add_attachments_section(story.attachments)
        
        self._add_appendix(ai_analysis)
        
        self.doc.save(output_path)
        print(f"Enhanced document saved: {output_path}")
    
    def _setup_enhanced_styles(self):
        """Set up enhanced document styles"""
        if self.styles_created:
            return
            
        styles = self.doc.styles
        
        # Cover page title
        cover_title = styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
        cover_title.font.size = Pt(24)
        cover_title.font.bold = True
        cover_title.font.color.rgb = RGBColor(0, 51, 102)
        cover_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cover_title.paragraph_format.space_after = Pt(12)
        
        # Section headers
        section_header = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        section_header.font.size = Pt(16)
        section_header.font.bold = True
        section_header.font.color.rgb = RGBColor(0, 51, 102)
        section_header.paragraph_format.space_before = Pt(12)
        section_header.paragraph_format.space_after = Pt(6)
        
        # Risk level styles
        high_risk = styles.add_style('HighRisk', WD_STYLE_TYPE.PARAGRAPH)
        high_risk.font.bold = True
        high_risk.font.color.rgb = RGBColor(204, 0, 0)
        
        medium_risk = styles.add_style('MediumRisk', WD_STYLE_TYPE.PARAGRAPH)
        medium_risk.font.bold = True
        medium_risk.font.color.rgb = RGBColor(255, 102, 0)
        
        low_risk = styles.add_style('LowRisk', WD_STYLE_TYPE.PARAGRAPH)
        low_risk.font.bold = True
        low_risk.font.color.rgb = RGBColor(0, 153, 0)
        
        self.styles_created = True
    
    def _add_cover_page(self, story):
        """Add professional cover page"""
        # Title
        title = self.doc.add_paragraph("SECURITY & FRAUD ANALYSIS REPORT", style='CoverTitle')
        
        # Story information
        self.doc.add_paragraph()
        story_info = self.doc.add_paragraph()
        story_info.add_run("Story: ").bold = True
        story_info.add_run(f"{story.key}")
        story_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        summary_info = self.doc.add_paragraph()
        summary_info.add_run("Summary: ").bold = True
        summary_info.add_run(story.summary)
        summary_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add some spacing
        for _ in range(10):
            self.doc.add_paragraph()
        
        # Document metadata
        meta_table = self.doc.add_table(rows=4, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        meta_data = [
            ("Generated On:", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ("Story Type:", story.story_type),
            ("Priority:", story.priority),
            ("Status:", story.status)
        ]
        
        for i, (key, value) in enumerate(meta_data):
            meta_table.cell(i, 0).text = key
            meta_table.cell(i, 1).text = str(value)
            meta_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Page break
        self.doc.add_page_break()
    
    def _add_executive_summary(self, story, ai_analysis: Dict[str, Any]):
        """Add executive summary section"""
        self.doc.add_paragraph("EXECUTIVE SUMMARY", style='SectionHeader')
        
        # Risk overview
        risk_assessment = ai_analysis.get('risk_assessment', {})
        risk_level = risk_assessment.get('level', 'Medium')
        risk_score = risk_assessment.get('score', 50)
        
        summary_para = self.doc.add_paragraph()
        summary_para.add_run("Overall Risk Level: ").bold = True
        
        risk_style = 'MediumRisk'
        if risk_level in ['Critical', 'High']:
            risk_style = 'HighRisk'
        elif risk_level == 'Low':
            risk_style = 'LowRisk'
        
        risk_para = self.doc.add_paragraph(f"{risk_level} ({risk_score}/100)", style=risk_style)
        
        # Key findings summary
        self.doc.add_paragraph()
        self.doc.add_paragraph("Key Findings:", style='Heading 3')
        
        fraud_count = len(ai_analysis.get('fraud_scenarios', []))
        security_count = len(ai_analysis.get('security_vulnerabilities', []))
        
        findings_para = self.doc.add_paragraph()
        findings_para.add_run(f"• {fraud_count} potential fraud scenarios identified\n")
        findings_para.add_run(f"• {security_count} security vulnerabilities found\n")
        findings_para.add_run(f"• {len(ai_analysis.get('mitigation_strategies', []))} mitigation strategies recommended")
        
        # Business impact
        business_impact = ai_analysis.get('business_impact', '')
        if business_impact:
            self.doc.add_paragraph()
            self.doc.add_paragraph("Business Impact:", style='Heading 3')
            self.doc.add_paragraph(business_impact)
        
        self.doc.add_page_break()
    
    def _add_risk_analysis(self, ai_analysis: Dict[str, Any]):
        """Add detailed risk analysis with visual elements"""
        self.doc.add_paragraph("RISK ANALYSIS", style='SectionHeader')
        
        risk_assessment = ai_analysis.get('risk_assessment', {})
        
        # Risk score visualization (text-based since we can't embed matplotlib easily)
        self.doc.add_paragraph("Risk Score Breakdown:", style='Heading 3')
        
        risk_score = risk_assessment.get('score', 50)
        pattern_matches = risk_assessment.get('pattern_matches', {})
        
        # Create a simple text-based risk visualization
        risk_bar = "█" * (risk_score // 5) + "░" * (20 - (risk_score // 5))
        risk_para = self.doc.add_paragraph(f"Risk Score: {risk_score}/100")
        risk_para.add_run(f"\n[{risk_bar}]")
        
        # Pattern analysis results
        self.doc.add_paragraph()
        self.doc.add_paragraph("Pattern Analysis Results:", style='Heading 3')
        
        security_patterns = pattern_matches.get('security_patterns', [])
        fraud_patterns = pattern_matches.get('fraud_patterns', [])
        
        if security_patterns:
            self.doc.add_paragraph("Security Patterns Detected:", style='Heading 4')
            for pattern in security_patterns:
                self.doc.add_paragraph(f"• {pattern}")
        
        if fraud_patterns:
            self.doc.add_paragraph("Fraud Patterns Detected:", style='Heading 4')
            for pattern in fraud_patterns:
                self.doc.add_paragraph(f"• {pattern}")
        
        if not security_patterns and not fraud_patterns:
            self.doc.add_paragraph("No specific risk patterns detected in the story content.")
    
    def _add_detailed_findings(self, ai_analysis: Dict[str, Any]):
        """Add detailed findings section"""
        self.doc.add_paragraph("DETAILED FINDINGS", style='SectionHeader')
        
        # Fraud scenarios
        self.doc.add_paragraph("Fraud Scenarios", style='Heading 3')
        fraud_scenarios = ai_analysis.get('fraud_scenarios', [])
        if fraud_scenarios:
            for i, scenario in enumerate(fraud_scenarios, 1):
                self.doc.add_paragraph(f"{i}. {scenario}")
        else:
            self.doc.add_paragraph("No specific fraud scenarios identified.")
        
        self.doc.add_paragraph()
        
        # Security vulnerabilities
        self.doc.add_paragraph("Security Vulnerabilities", style='Heading 3')
        security_vulns = ai_analysis.get('security_vulnerabilities', [])
        if security_vulns:
            for i, vuln in enumerate(security_vulns, 1):
                self.doc.add_paragraph(f"{i}. {vuln}")
        else:
            self.doc.add_paragraph("No specific security vulnerabilities identified.")
    
    def _add_recommendations(self, ai_analysis: Dict[str, Any]):
        """Add recommendations section"""
        self.doc.add_paragraph("RECOMMENDATIONS", style='SectionHeader')
        
        # Mitigation strategies
        self.doc.add_paragraph("Mitigation Strategies", style='Heading 3')
        strategies = ai_analysis.get('mitigation_strategies', [])
        if strategies:
            for i, strategy in enumerate(strategies, 1):
                self.doc.add_paragraph(f"{i}. {strategy}")
        
        self.doc.add_paragraph()
        
        # Testing recommendations
        self.doc.add_paragraph("Testing Recommendations", style='Heading 3')
        testing_recs = ai_analysis.get('testing_recommendations', [])
        if testing_recs:
            for i, rec in enumerate(testing_recs, 1):
                self.doc.add_paragraph(f"{i}. {rec}")
        else:
            self.doc.add_paragraph("Standard security testing protocols should be applied.")
    
    def _add_compliance_section(self, ai_analysis: Dict[str, Any]):
        """Add compliance considerations section"""
        self.doc.add_paragraph("COMPLIANCE CONSIDERATIONS", style='SectionHeader')
        
        compliance_items = ai_analysis.get('compliance_impact', [])
        if compliance_items:
            for i, item in enumerate(compliance_items, 1):
                self.doc.add_paragraph(f"{i}. {item}")
        else:
            self.doc.add_paragraph("No specific compliance requirements identified. Review applicable regulations based on your industry and data handling practices.")
    
    def _add_appendix(self, ai_analysis: Dict[str, Any]):
        """Add appendix with technical details"""
        self.doc.add_page_break()
        self.doc.add_paragraph("APPENDIX", style='SectionHeader')
        
        # Analysis metadata
        metadata = ai_analysis.get('analysis_metadata', {})
        if metadata:
            self.doc.add_paragraph("Analysis Metadata", style='Heading 3')
            
            meta_table = self.doc.add_table(rows=len(metadata), cols=2)
            meta_table.style = 'Table Grid'
            
            for i, (key, value) in enumerate(metadata.items()):
                meta_table.cell(i, 0).text = key.replace('_', ' ').title()
                meta_table.cell(i, 1).text = str(value)
                meta_table.cell(i, 0).paragraphs[0].runs[0].bold = True
```

## README Documentation
```markdown
# Jira Story to Word Document Generator

An end-to-end solution for generating comprehensive Word documents from Jira stories with AI-powered fraud and security analysis.

## Features

- **Jira Integration**: Seamlessly connect to Jira and extract story details
- **AI-Powered Analysis**: Use OpenAI GPT models for fraud and security analysis
- **Professional Documents**: Generate well-formatted Word documents
- **Attachment Processing**: Download and analyze Jira attachments
- **Pattern Matching**: Built-in security and fraud pattern detection
- **Batch Processing**: Process multiple stories efficiently
- **Configurable**: Flexible configuration options

## Quick Start

1. **Clone and Setup**
   ```bash
   git clone <repository-url>
   cd jira-word-generator
   python scripts/setup_env.py
   ```

2. **Configure Environment**
   ```bash
   cp .env.example .env
   # Edit .env with your credentials
   ```

3. **Validate Configuration**
   ```bash
   python scripts/validate_config.py
   ```

4. **Process a Story**
   ```bash
   python src/main.py PROJ-123
   ```

## Installation

### Prerequisites
- Python 3.8+
- Jira account with API access
- OpenAI API key

### Setup
```bash
# Install dependencies
pip install -r requirements.txt

# Create necessary directories
mkdir -p output/{documents,attachments} logs templates

# Copy configuration template
cp config/config.yaml.example config/config.yaml
```

## Configuration

### Environment Variables (.env)
```env
JIRA_SERVER=https://your-company.atlassian.net
JIRA_USERNAME=your-email@company.com
JIRA_API_TOKEN=your-jira-api-token
OPENAI_API_KEY=your-openai-api-key
```

### YAML Configuration (config/config.yaml)
See `config/config.yaml.example` for detailed configuration options.

## Usage

### Single Story Processing
```python
from src.main import JiraWordGenerator
from config.settings import load_config

config = load_config()
generator = JiraWordGenerator(config.dict())
output_path = generator.process_story("PROJ-123")
```

### Batch Processing
```bash
# Create CSV with story keys
echo "PROJ-123\nPROJ-124\nPROJ-125" > stories.csv

# Process batch
python scripts/batch_process.py stories.csv
```

### API Usage
```python
from src.jira_connector import JiraConnector
from src.ai_analyzer import EnhancedAIAnalyzer
from src.word_generator import EnhancedWordGenerator

# Initialize components
jira = JiraConnector(server, username, token)
analyzer = EnhancedAIAnalyzer(openai_key)
generator = EnhancedWordGenerator()

# Process story
story = jira.get_story("PROJ-123")
analysis = analyzer.analyze_fraud_security_cases(story)
generator.create_document(story, analysis, "output.docx")
```

## Document Structure

Generated documents include:

1. **Cover Page** - Story summary and metadata
2. **Executive Summary** - Key findings and risk overview
3. **Story Details** - Complete Jira story information
4. **Risk Analysis** - Pattern matching and AI analysis results
5. **Detailed Findings** - Fraud scenarios and security vulnerabilities
6. **Recommendations** - Mitigation strategies and testing guidance
7. **Compliance Considerations** - Regulatory requirements
8. **Attachments** - Information about story attachments
9. **Appendix** - Technical analysis metadata

## Security Analysis Features

### Pattern Detection
- Authentication bypass patterns
- SQL injection indicators
- XSS vulnerability markers
- Data exposure risks
- Privilege escalation patterns

### Fraud Detection
- Payment fraud scenarios
- Identity theft risks
- Data manipulation threats
- Social engineering vectors

### AI Analysis
- Contextual risk assessment
- Business impact evaluation
- Compliance requirement identification
- Tailored mitigation strategies

## Testing

```bash
# Run all tests
pytest

# Run with coverage
pytest --cov=src

# Run specific test file
pytest tests/test_jira_connector.py
```

## Contributing

1. Fork the repository
2. Create a feature branch
3. Make changes with tests
4. Submit a pull request

## License

MIT License - see LICENSE file for details.

## Support

For issues and questions:
- Check the documentation in `/docs`
- Review configuration examples
- Run validation scripts
- Create GitHub issues for bugs

## Changelog

### v1.0.0
- Initial release
- Basic Jira integration
- AI-powered analysis
- Word document generation
- Batch processing support
```



















# Test Files and Additional Utilities

## tests/test_jira_connector.py
import pytest
import json
from unittest.mock import Mock, patch
from src.jira_connector import JiraConnector, JiraStory

class TestJiraConnector:
    
    @pytest.fixture
    def mock_jira_issue(self):
        """Mock Jira issue object"""
        issue = Mock()
        issue.key = "TEST-123"
        issue.fields.summary = "Test Story Summary"
        issue.fields.description = "Test description with Acceptance Criteria: Given user logs in When they click submit Then they see success"
        issue.fields.labels = ["security", "api"]
        issue.fields.attachment = []
        issue.fields.issuetype = Mock()
        issue.fields.issuetype.__str__ = lambda x: "Story"
        issue.fields.priority = Mock()
        issue.fields.priority.__str__ = lambda x: "High"
        issue.fields.status = Mock()
        issue.fields.status.__str__ = lambda x: "In Progress"
        issue.fields.assignee = Mock()
        issue.fields.assignee.__str__ = lambda x: "john.doe@company.com"
        issue.fields.reporter = Mock()
        issue.fields.reporter.__str__ = lambda x: "jane.smith@company.com"
        issue.fields.created = "2024-01-01T10:00:00.000+0000"
        issue.fields.updated = "2024-01-02T10:00:00.000+0000"
        return issue
    
    @pytest.fixture
    def jira_connector(self):
        """Create JiraConnector instance with mocked JIRA"""
        with patch('src.jira_connector.JIRA') as mock_jira:
            connector = JiraConnector("https://test.atlassian.net", "user", "token")
            connector.jira = mock_jira.return_value
            return connector
    
    def test_get_story_success(self, jira_connector, mock_jira_issue):
        """Test successful story retrieval"""
        jira_connector.jira.issue.return_value = mock_jira_issue
        
        story = jira_connector.get_story("TEST-123")
        
        assert isinstance(story, JiraStory)
        assert story.key == "TEST-123"
        assert story.summary == "Test Story Summary"
        assert "Given user logs in" in story.acceptance_criteria
        assert "security" in story.labels
    
    def test_extract_acceptance_criteria(self, jira_connector, mock_jira_issue):
        """Test acceptance criteria extraction"""
        # Test extraction from description
        ac = jira_connector._extract_acceptance_criteria(mock_jira_issue)
        assert "Given user logs in" in ac
        
        # Test with no AC in description
        mock_jira_issue.fields.description = "Simple description without AC"
        ac = jira_connector._extract_acceptance_criteria(mock_jira_issue)
        assert ac == ""
    
    def test_download_attachment_success(self, jira_connector):
        """Test successful attachment download"""
        mock_response = Mock()
        mock_response.status_code = 200
        mock_response.content = b"test file content"
        
        jira_connector.jira._session.get.return_value = mock_response
        
        with patch('builtins.open', create=True) as mock_open:
            result = jira_connector.download_attachment("http://test.com/file", "test.txt")
            assert result is True
            mock_open.assert_called_once()

## tests/test_ai_analyzer.py
import pytest
import json
from unittest.mock import Mock, patch
from src.ai_analyzer import EnhancedAIAnalyzer, SecurityPattern
from src.jira_connector import JiraStory

class TestEnhancedAIAnalyzer:
    
    @pytest.fixture
    def sample_story(self):
        """Create sample JiraStory for testing"""
        return JiraStory(
            key="TEST-123",
            summary="User login authentication system",
            description="Implement secure user authentication with password validation",
            acceptance_criteria="Given user enters valid credentials When they submit Then they are authenticated",
            labels=["security", "authentication"],
            attachments=[],
            story_type="Story",
            priority="High",
            status="To Do",
            assignee="developer@company.com",
            reporter="pm@company.com",
            created="2024-01-01T10:00:00.000+0000",
            updated="2024-01-01T10:00:00.000+0000"
        )
    
    @pytest.fixture
    def ai_analyzer(self):
        """Create EnhancedAIAnalyzer instance"""
        return EnhancedAIAnalyzer("test-api-key")
    
    def test_load_security_patterns(self, ai_analyzer):
        """Test security patterns loading"""
        patterns = ai_analyzer._load_security_patterns()
        assert len(patterns) > 0
        assert any(p.name == "Authentication Bypass" for p in patterns)
        assert any(p.name == "SQL Injection" for p in patterns)
    
    def test_load_fraud_patterns(self, ai_analyzer):
        """Test fraud patterns loading"""
        patterns = ai_analyzer._load_fraud_patterns()
        assert len(patterns) > 0
        assert any(p.name == "Payment Fraud" for p in patterns)
        assert any(p.name == "Identity Fraud" for p in patterns)
    
    def test_analyze_patterns(self, ai_analyzer, sample_story):
        """Test pattern analysis"""
        content = f"{sample_story.summary} {sample_story.description}"
        results = ai_analyzer.analyze_patterns(content)
        
        assert "security_patterns" in results
        assert "fraud_patterns" in results
        assert len(results["security_patterns"]) > 0  # Should match authentication patterns
    
    def test_calculate_risk_score(self, ai_analyzer):
        """Test risk score calculation"""
        pattern_analysis = {
            "security_patterns": ["Auth pattern", "SQL pattern"],
            "fraud_patterns": ["Payment pattern"]
        }
        ai_analysis = {
            "risk_assessment": {"level": "High"}
        }
        
        score = ai_analyzer._calculate_risk_score(pattern_analysis, ai_analysis)
        assert 0 <= score <= 100
        assert score > 50  # Should be high due to multiple patterns and High risk level
    
    @patch('openai.ChatCompletion.create')
    def test_perform_ai_analysis_success(self, mock_openai, ai_analyzer, sample_story):
        """Test successful AI analysis"""
        mock_response = Mock()
        mock_response.choices = [Mock()]
        mock_response.choices[0].message.content = '''
        {
            "fraud_scenarios": ["Credential theft", "Account takeover"],
            "security_vulnerabilities": ["Weak password policy", "No MFA"],
            "risk_assessment": {"level": "High", "score": 80, "justification": "Multiple auth risks"},
            "mitigation_strategies": ["Implement MFA", "Strong password policy"],
            "testing_recommendations": ["Penetration testing", "Security code review"],
            "compliance_impact": ["GDPR compliance required"],
            "business_impact": "Potential data breach and customer trust loss"
        }
        '''
        mock_openai.return_value = mock_response
        
        content = f"{sample_story.summary} {sample_story.description}"
        pattern_analysis = {"security_patterns": [], "fraud_patterns": []}
        
        result = ai_analyzer._perform_ai_analysis(content, pattern_analysis)
        
        assert "fraud_scenarios" in result
        assert len(result["fraud_scenarios"]) == 2
        assert result["risk_assessment"]["level"] == "High"
    
    @patch('openai.ChatCompletion.create')
    def test_perform_ai_analysis_failure(self, mock_openai, ai_analyzer, sample_story):
        """Test AI analysis failure handling"""
        mock_openai.side_effect = Exception("API Error")
        
        content = f"{sample_story.summary} {sample_story.description}"
        pattern_analysis = {"security_patterns": [], "fraud_patterns": []}
        
        result = ai_analyzer._perform_ai_analysis(content, pattern_analysis)
        
        # Should return fallback analysis
        assert result["fraud_scenarios"] == ["Manual review required - AI analysis unavailable"]

## tests/test_word_generator.py
import pytest
import tempfile
import os
from unittest.mock import Mock, patch
from src.word_generator import EnhancedWordGenerator
from src.jira_connector import JiraStory

class TestEnhancedWordGenerator:
    
    @pytest.fixture
    def sample_story(self):
        """Create sample JiraStory for testing"""
        return JiraStory(
            key="TEST-123",
            summary="Test Story Summary",
            description="Test description",
            acceptance_criteria="Given When Then",
            labels=["test", "security"],
            attachments=[{
                "filename": "test.pdf",
                "size": 1024,
                "created": "2024-01-01T10:00:00.000+0000",
                "content_url": "http://test.com/file",
                "mime_type": "application/pdf"
            }],
            story_type="Story",
            priority="High",
            status="To Do",
            assignee="developer@company.com",
            reporter="pm@company.com",
            created="2024-01-01T10:00:00.000+0000",
            updated="2024-01-01T10:00:00.000+0000"
        )
    
    @pytest.fixture
    def sample_analysis(self):
        """Create sample AI analysis result"""
        return {
            "fraud_scenarios": ["Test fraud scenario"],
            "security_vulnerabilities": ["Test security vulnerability"],
            "risk_assessment": {
                "level": "High",
                "score": 80,
                "justification": "Multiple risks identified",
                "pattern_matches": {
                    "security_patterns": ["Auth pattern"],
                    "fraud_patterns": ["Payment pattern"]
                }
            },
            "mitigation_strategies": ["Implement security controls"],
            "testing_recommendations": ["Conduct security testing"],
            "compliance_impact": ["GDPR compliance required"],
            "business_impact": "Potential security breach",
            "analysis_metadata": {
                "model_used": "gpt-4",
                "pattern_matches_count": 2,
                "analysis_timestamp": "2024-01-01T10:00:00"
            }
        }
    
    @pytest.fixture
    def word_generator(self):
        """Create EnhancedWordGenerator instance"""
        return EnhancedWordGenerator()
    
    def test_create_document(self, word_generator, sample_story, sample_analysis):
        """Test document creation"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            try:
                word_generator.create_document(sample_story, sample_analysis, tmp_file.name)
                assert os.path.exists(tmp_file.name)
                assert os.path.getsize(tmp_file.name) > 0
            finally:
                os.unlink(tmp_file.name)
    
    def test_setup_enhanced_styles(self, word_generator):
        """Test style setup"""
        from docx import Document
        word_generator.doc = Document()
        word_generator._setup_enhanced_styles()
        
        assert word_generator.styles_created
        assert 'CoverTitle' in [style.name for style in word_generator.doc.styles]
    
    def test_add_cover_page(self, word_generator, sample_story):
        """Test cover page creation"""
        from docx import Document
        word_generator.doc = Document()
        word_generator._setup_enhanced_styles()
        word_generator._add_cover_page(sample_story)
        
        # Check that content was added
        assert len(word_generator.doc.paragraphs) > 0
        assert any("SECURITY & FRAUD ANALYSIS REPORT" in p.text for p in word_generator.doc.paragraphs)

## tests/fixtures/sample_story.json
{
    "key": "PROJ-123",
    "summary": "Implement user authentication system",
    "description": "Create a secure authentication system that allows users to log in with email and password. The system should include password validation, rate limiting, and session management.",
    "acceptance_criteria": "Given a user with valid credentials\nWhen they attempt to log in\nThen they should be authenticated and redirected to the dashboard\n\nGiven a user with invalid credentials\nWhen they attempt to log in\nThen they should see an error message and remain on the login page",
    "labels": ["authentication", "security", "backend"],
    "attachments": [
        {
            "filename": "auth_flow_diagram.png",
            "size": 245760,
            "created": "2024-01-01T10:00:00.000+0000",
            "content_url": "https://company.atlassian.net/secure/attachment/123456/auth_flow_diagram.png",
            "mime_type": "image/png"
        }
    ],
    "story_type": "Story",
    "priority": "High",
    "status": "In Progress",
    "assignee": "john.doe@company.com",
    "reporter": "jane.smith@company.com",
    "created": "2024-01-01T09:00:00.000+0000",
    "updated": "2024-01-01T15:30:00.000+0000"
}

## tests/fixtures/sample_analysis.json
{
    "fraud_scenarios": [
        "Credential stuffing attacks using compromised email/password combinations",
        "Account takeover through session hijacking",
        "Brute force attacks against user accounts",
        "Social engineering to obtain user credentials"
    ],
    "security_vulnerabilities": [
        "Insufficient password complexity requirements",
        "Lack of multi-factor authentication",
        "Session tokens without proper expiration",
        "No rate limiting on login attempts",
        "Potential for timing attacks during authentication"
    ],
    "risk_assessment": {
        "level": "High",
        "score": 85,
        "justification": "Authentication systems are high-value targets for attackers. Multiple attack vectors exist including credential theft, session hijacking, and brute force attacks. The absence of MFA significantly increases risk.",
        "pattern_matches": {
            "security_patterns": [
                "Authentication Bypass - Potential authentication bypass vulnerabilities",
                "Privilege Escalation - Privilege escalation vulnerabilities"
            ],
            "fraud_patterns": [
                "Identity Fraud - Identity theft and impersonation risks"
            ]
        }
    },
    "mitigation_strategies": [
        "Implement multi-factor authentication (MFA)",
        "Enforce strong password policies with complexity requirements",
        "Add rate limiting and account lockout mechanisms",
        "Implement secure session management with proper token expiration",
        "Use secure password storage with proper hashing (bcrypt/Argon2)",
        "Add login attempt monitoring and alerting",
        "Implement CAPTCHA for suspicious login patterns"
    ],
    "testing_recommendations": [
        "Conduct penetration testing focusing on authentication bypass",
        "Perform automated security scanning for common vulnerabilities",
        "Test for timing attacks and information disclosure",
        "Validate rate limiting and account lockout functionality",
        "Review session management implementation",
        "Test password reset functionality for security flaws"
    ],
    "compliance_impact": [
        "GDPR - Ensure proper consent and data protection for user credentials",
        "PCI DSS - If processing payments, secure authentication is required",
        "SOX - Implement proper access controls for financial systems",
        "HIPAA - If handling health data, strong authentication is mandatory"
    ],
    "business_impact": "A successful attack could result in unauthorized access to user accounts, data breaches, financial losses, regulatory fines, and severe damage to company reputation and customer trust.",
    "analysis_metadata": {
        "model_used": "gpt-4",
        "pattern_matches_count": 3,
        "analysis_timestamp": "2024-01-01T10:00:00.000000"
    }
}

## src/utils/file_handler.py
"""
File handling utilities for processing attachments and documents
"""

import os
import mimetypes
import magic
from typing import List, Dict, Any, Optional
import pandas as pd
import json
from pathlib import Path

class FileHandler:
    """Handle various file types and extract content for analysis"""
    
    SUPPORTED_TEXT_TYPES = [
        'text/plain', 'text/csv', 'text/markdown',
        'application/json', 'application/xml'
    ]
    
    SUPPORTED_DOC_TYPES = [
        'application/pdf', 'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/vnd.ms-excel',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    ]
    
    def __init__(self):
        self.mime = magic.Magic(mime=True)
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file MIME type"""
        try:
            return self.mime.from_file(file_path)
        except Exception:
            # Fallback to extension-based detection
            mime_type, _ = mimetypes.guess_type(file_path)
            return mime_type or 'application/octet-stream'
    
    def extract_text_content(self, file_path: str) -> Optional[str]:
        """Extract text content from various file types"""
        if not os.path.exists(file_path):
            return None
        
        mime_type = self.detect_file_type(file_path)
        
        try:
            if mime_type in self.SUPPORTED_TEXT_TYPES:
                return self._extract_text_file(file_path)
            elif mime_type == 'application/pdf':
                return self._extract_pdf_content(file_path)
            elif 'excel' in mime_type or 'spreadsheet' in mime_type:
                return self._extract_excel_content(file_path)
            elif 'word' in mime_type or 'document' in mime_type:
                return self._extract_word_content(file_path)
            else:
                return f"Binary file: {os.path.basename(file_path)} ({mime_type})"
        except Exception as e:
            return f"Error reading file {os.path.basename(file_path)}: {str(e)}"
    
    def _extract_text_file(self, file_path: str) -> str:
        """Extract content from text files"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        return f"Could not decode text file: {os.path.basename(file_path)}"
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except ImportError:
            return "PDF processing requires PyPDF2 library"
        except Exception as e:
            return f"Error extracting PDF content: {str(e)}"
    
    def _extract_excel_content(self, file_path: str) -> str:
        """Extract content from Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                content.append(f"Sheet: {sheet_name}")
                content.append(df.to_string())
                content.append("")
            
            return "\n".join(content)
        except Exception as e:
            return f"Error extracting Excel content: {str(e)}"
    
    def _extract_word_content(self, file_path: str) -> str:
        """Extract content from Word documents"""
        try:
            from docx import Document
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    content.append(" | ".join(row_text))
            
            return "\n".join(content)
        except Exception as e:
            return f"Error extracting Word content: {str(e)}"
    
    def get_file_summary(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file summary"""
        if not os.path.exists(file_path):
            return {"error": "File not found"}
        
        stat = os.stat(file_path)
        mime_type = self.detect_file_type(file_path)
        
        return {
            "filename": os.path.basename(file_path),
            "size": stat.st_size,
            "mime_type": mime_type,
            "is_text_extractable": mime_type in (self.SUPPORTED_TEXT_TYPES + self.SUPPORTED_DOC_TYPES),
            "modified_time": stat.st_mtime,
            "path": file_path
        }

## src/utils/validators.py
"""
Validation utilities for configuration and input data
"""

import re
import requests
from typing import Dict, List, Any, Optional
from urllib.parse import urlparse

class ConfigValidator:
    """Validate configuration settings"""
    
    @staticmethod
    def validate_jira_config(server: str, username: str, token: str) -> Dict[str, Any]:
        """Validate Jira configuration"""
        result = {"valid": True, "errors": []}
        
        # Validate server URL
        if not server:
            result["errors"].append("Jira server URL is required")
        elif not server.startswith(('http://', 'https://')):
            result["errors"].append("Jira server URL must start with http:// or https://")
        else:
            parsed = urlparse(server)
            if not parsed.netloc:
                result["errors"].append("Invalid Jira server URL format")
        
        # Validate username
        if not username:
            result["errors"].append("Jira username is required")
        elif not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}, username):
            result["errors"].append("Jira username should be a valid email address")
        
        # Validate token
        if not token:
            result["errors"].append("Jira API token is required")
        elif len(token) < 10:
            result["errors"].append("Jira API token appears to be too short")
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    @staticmethod
    def validate_openai_config(api_key: str, model: str = "gpt-4") -> Dict[str, Any]:
        """Validate OpenAI configuration"""
        result = {"valid": True, "errors": []}
        
        # Validate API key
        if not api_key:
            result["errors"].append("OpenAI API key is required")
        elif not api_key.startswith('sk-'):
            result["errors"].append("OpenAI API key should start with 'sk-'")
        elif len(api_key) < 40:
            result["errors"].append("OpenAI API key appears to be too short")
        
        # Validate model
        valid_models = ['gpt-4', 'gpt-4-turbo', 'gpt-3.5-turbo']
        if model not in valid_models:
            result["errors"].append(f"Model should be one of: {', '.join(valid_models)}")
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    @staticmethod
    def validate_story_key(story_key: str) -> Dict[str, Any]:
        """Validate Jira story key format"""
        result = {"valid": True, "errors": []}
        
        if not story_key:
            result["errors"].append("Story key is required")
        elif not re.match(r'^[A-Z]+-\d+, story_key):
            result["errors"].append("Story key should follow format: PROJECT-123")
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    @staticmethod
    def validate_output_path(path: str) -> Dict[str, Any]:
        """Validate output path"""
        result = {"valid": True, "errors": []}
        
        if not path:
            result["errors"].append("Output path is required")
        elif not path.endswith('.docx'):
            result["errors"].append("Output path should end with .docx")
        
        # Check if directory exists
        import os
        directory = os.path.dirname(path)
        if directory and not os.path.exists(directory):
            try:
                os.makedirs(directory, exist_ok=True)
            except Exception as e:
                result["errors"].append(f"Cannot create output directory: {str(e)}")
        
        result["valid"] = len(result["errors"]) == 0
        return result

class DataValidator:
    """Validate data integrity and completeness"""
    
    @staticmethod
    def validate_story_data(story_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate story data completeness"""
        result = {"valid": True, "warnings": [], "errors": []}
        
        required_fields = ['key', 'summary', 'story_type', 'status']
        for field in required_fields:
            if not story_data.get(field):
                result["errors"].append(f"Missing required field: {field}")
        
        # Check for recommended fields
        recommended_fields = ['description', 'acceptance_criteria', 'assignee']
        for field in recommended_fields:
            if not story_data.get(field):
                result["warnings"].append(f"Missing recommended field: {field}")
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    @staticmethod
    def validate_analysis_result(analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Validate AI analysis result structure"""
        result = {"valid": True, "errors": []}
        
        required_keys = [
            'fraud_scenarios', 'security_vulnerabilities', 'risk_assessment',
            'mitigation_strategies', 'testing_recommendations'
        ]
        
        for key in required_keys:
            if key not in analysis:
                result["errors"].append(f"Missing analysis section: {key}")
        
        # Validate risk assessment structure
        if 'risk_assessment' in analysis:
            risk_assessment = analysis['risk_assessment']
            if not isinstance(risk_assessment, dict):
                result["errors"].append("Risk assessment should be a dictionary")
            elif 'level' not in risk_assessment:
                result["errors"].append("Risk assessment missing level")
        
        result["valid"] = len(result["errors"]) == 0
        return result

## Dockerfile
FROM python:3.11-slim

WORKDIR /app

# Install system dependencies
RUN apt-get update && apt-get install -y \
    libmagic1 \
    libmagic-dev \
    && rm -rf /var/lib/apt/lists/*

# Copy requirements and install Python dependencies
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

# Copy application code
COPY src/ ./src/
COPY config/ ./config/
COPY templates/ ./templates/
COPY scripts/ ./scripts/

# Create necessary directories
RUN mkdir -p output/documents output/attachments logs

# Set environment variables
ENV PYTHONPATH=/app/src

# Expose port for potential web interface
EXPOSE 8000

# Default command
CMD ["python", "src/main.py", "--help"]

## docker-compose.yml
version: '3.8'

services:
  jira-word-generator:
    build: .
    volumes:
      - ./output:/app/output
      - ./logs:/app/logs
      - ./config:/app/config
    environment:
      - JIRA_SERVER=${JIRA_SERVER}
      - JIRA_USERNAME=${JIRA_USERNAME}
      - JIRA_API_TOKEN=${JIRA_API_TOKEN}
      - OPENAI_API_KEY=${OPENAI_API_KEY}
    env_file:
      - .env
    command: python src/main.py --config config/config.yaml

## Makefile
.PHONY: install test lint format clean build docker-build docker-run

# Installation
install:
	python -m venv venv
	. venv/bin/activate && pip install -r requirements.txt
	python scripts/setup_env.py

# Testing
test:
	pytest tests/ -v

test-coverage:
	pytest tests/ --cov=src --cov-report=html --cov-report=term

# Code quality
lint:
	flake8 src/ tests/
	mypy src/

format:
	black src/ tests/ scripts/
	isort src/ tests/ scripts/

# Validation
validate:
	python scripts/validate_config.py

# Build
build:
	python setup.py sdist bdist_wheel

# Docker
docker-build:
	docker build -t jira-word-generator .

docker-run:
	docker-compose up

# Cleanup
clean:
	rm -rf build/ dist/ *.egg-info/
	find . -type d -name __pycache__ -delete
	find . -type f -name "*.pyc" -delete

# Development
dev-setup: install
	pre-commit install

# Documentation
docs:
	sphinx-build -b html docs/ docs/_build/

# Release
release: test lint build
	twine upload dist/*











